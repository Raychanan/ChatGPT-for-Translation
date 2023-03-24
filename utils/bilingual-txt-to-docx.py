from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from langdetect import detect
from docx.oxml.ns import qn

def is_chinese(text):
    detected_lang = detect(text)
    return detected_lang == 'zh-cn' or detected_lang == 'zh-tw'

def create_bilingual_docx(txt_file, output_docx):
    with open(txt_file, 'r', encoding='utf-8') as f:
        paragraphs = f.readlines()

    document = Document()
    document.styles['Normal'].font.name = u'Source Han Serif SC'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Source Han Serif SC')
    document.styles['Normal'].font.size = Pt(7)

    # Create English style
    english_style = document.styles.add_style('EnglishStyle', WD_STYLE_TYPE.PARAGRAPH)
    english_font = english_style.font
    english_style.font.name = "Times New Roman"
    english_font.size = Pt(11)
    english_style.paragraph_format.line_spacing = 1.2

    # Apply styles to paragraphs
    for paragraph_text in paragraphs:
        paragraph_text = paragraph_text.strip()
        if is_chinese(paragraph_text):
            paragraph = document.add_paragraph(paragraph_text, style='Normal')
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(0)
        else:
            paragraph = document.add_paragraph(paragraph_text, style='EnglishStyle')
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(0)



    document.save(output_docx)

# Example usage:
create_bilingual_docx('../input_bilingual.txt', 'output.docx')
